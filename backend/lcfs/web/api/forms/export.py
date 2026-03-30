"""Base PDF and DOCX exporters """

import io
import os
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel
from starlette.responses import StreamingResponse

from lcfs.web.api.forms.schema import DeclarationExportRequest

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")
_JINJA_ENV    = Environment(loader=FileSystemLoader(_TEMPLATE_DIR), autoescape=True)

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ── Base exporters ────────────────────────────────────────────────────────────


class BasePdfExporter:
    """Renders a Jinja2 template via WeasyPrint. Subclasses supply the template and context."""

    template_file: str

    def build_context(self, payload: BaseModel, form_name: Optional[str]) -> dict:
        raise NotImplementedError

    def safe_filename(self, payload: BaseModel, ext: str) -> str:
        raise NotImplementedError

    async def export(
        self, payload: BaseModel, form_name: Optional[str] = None
    ) -> StreamingResponse:
        from weasyprint import HTML  # heavy import — defer until needed

        ctx      = self.build_context(payload, form_name)
        html_str = _JINJA_ENV.get_template(self.template_file).render(**ctx)
        pdf      = HTML(string=html_str, base_url=_TEMPLATE_DIR).write_pdf()
        filename = self.safe_filename(payload, "pdf")

        return StreamingResponse(
            io.BytesIO(pdf),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


class BaseDocxExporter:
    """Builds a Word document with python-docx. Subclasses implement ``_build_doc``."""

    def build_context(self, payload: BaseModel, form_name: Optional[str]) -> dict:
        raise NotImplementedError

    def safe_filename(self, payload: BaseModel, ext: str) -> str:
        raise NotImplementedError

    def _build_doc(self, ctx: dict) -> Document:
        raise NotImplementedError

    async def export(
        self, payload: BaseModel, form_name: Optional[str] = None
    ) -> StreamingResponse:
        ctx    = self.build_context(payload, form_name)
        doc    = self._build_doc(ctx)
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        filename = self.safe_filename(payload, "docx")

        return StreamingResponse(
            stream,
            media_type=_DOCX_MEDIA_TYPE,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


# ── Fuel Supplier Declaration exporters ───────────────────────────────────────
# Shared helpers (used by both PDF and DOCX for this form)

_DECLARATION_DEFAULT_NAME = "Fuel Supplier Declaration"


def _declaration_context(
    payload: DeclarationExportRequest, form_name: Optional[str]
) -> dict:
    """Shared context for both the PDF template and DOCX builder."""
    now = datetime.now(timezone.utc)
    return {
        "form_name":         form_name or _DECLARATION_DEFAULT_NAME,
        "organization_name": payload.organization_name,
        "reporting_period":  payload.reporting_period,
        "declaration_type":  payload.declaration_type,
        "contact_name":      payload.contact_name,
        "contact_email":     payload.contact_email,
        "fuel_type":         payload.fuel_type,
        "quantity":          f"{payload.quantity:,.2f}",
        "units":             payload.units,
        "notes":             payload.notes or "",
        "certified":         payload.certified,
        "certified_date":    now.strftime("%B %d, %Y"),
        "generated":         now.strftime("%B %d, %Y at %H:%M UTC"),
        "generated_date":    now.strftime("%Y-%m-%d"),
    }


def _declaration_filename(payload: DeclarationExportRequest, ext: str) -> str:
    safe = "".join(c if c.isalnum() else "_" for c in payload.organization_name)
    return f"LCFS-Declaration-{safe}-{payload.reporting_period}.{ext}"


class FuelSupplierDeclarationPdfExporter(BasePdfExporter):
    template_file = "fuel-supplier-declaration.html"

    def build_context(
        self, payload: DeclarationExportRequest, form_name: Optional[str]
    ) -> dict:
        return _declaration_context(payload, form_name)

    def safe_filename(self, payload: DeclarationExportRequest, ext: str) -> str:
        return _declaration_filename(payload, ext)


class FuelSupplierDeclarationDocxExporter(BaseDocxExporter):
    _FONT = "Arial"
    _CERT = (
        "I, the undersigned authorized representative, hereby certify that the "
        "information provided in this declaration is true, accurate, and complete "
        "to the best of my knowledge. I understand that providing false or "
        "misleading information may result in penalties under the Low Carbon Fuels Act."
    )

    def build_context(
        self, payload: DeclarationExportRequest, form_name: Optional[str]
    ) -> dict:
        return _declaration_context(payload, form_name)

    def safe_filename(self, payload: DeclarationExportRequest, ext: str) -> str:
        return _declaration_filename(payload, ext)

    # ── Document assembly ─────────────────────────────────────────────────────

    def _build_doc(self, ctx: dict) -> Document:
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.5)

        self._p(doc, "Government of British Columbia \u2014 Low Carbon Fuels Act", size=9)
        self._p(doc, ctx["form_name"], bold=True, size=16)
        self._p(
            doc,
            f"Reporting Period: {ctx['reporting_period']}  \u00b7  Generated: {ctx['generated']}",
            size=8,
        )
        doc.add_paragraph()

        self._heading(doc, "Organization Information")
        self._info_table(doc, [
            ("Organization",     ctx["organization_name"]),
            ("Compliance Period", ctx["reporting_period"]),
            ("Declaration Type", ctx["declaration_type"]),
        ])

        self._heading(doc, "Contact")
        self._info_table(doc, [
            ("Name",  ctx["contact_name"]),
            ("Email", ctx["contact_email"]),
        ])

        self._heading(doc, "Fuel Supply")
        self._info_table(doc, [
            ("Fuel Type", ctx["fuel_type"]),
            ("Quantity",  f"{ctx['quantity']} {ctx['units']}"),
        ])

        if ctx["notes"]:
            self._heading(doc, "Additional Notes")
            p = doc.add_paragraph()
            r = p.add_run(ctx["notes"])
            r.font.size = Pt(9)
            r.font.name = self._FONT
            doc.add_paragraph()

        self._heading(doc, "Certification")
        self._cert_block(doc, ctx)

        self._page_footer(doc, ctx)
        return doc

    # ── Section builders ──────────────────────────────────────────────────────

    def _p(self, doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
        r = doc.add_paragraph().add_run(str(text))
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = self._FONT

    def _heading(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = self._FONT

    def _cell(
        self,
        cell,
        text: str,
        bold: bool = False,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> None:
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(str(text))
        r.bold = bold
        r.font.size = Pt(9)
        r.font.name = self._FONT

    def _info_table(self, doc: Document, rows: list[tuple[str, str]]) -> None:
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            self._cell(tbl.rows[i].cells[0], label, bold=True)
            self._cell(tbl.rows[i].cells[1], value)
        doc.add_paragraph()

    def _cert_block(self, doc: Document, ctx: dict) -> None:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]

        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(6)
        r1 = p1.add_run(self._CERT)
        r1.italic = True
        r1.font.size = Pt(9)
        r1.font.name = self._FONT

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(
            f"\u2611  Certified by {ctx['contact_name']} on {ctx['certified_date']}"
        )
        r2.bold = True
        r2.font.size = Pt(9)
        r2.font.name = self._FONT

        doc.add_paragraph()

        sig = doc.add_paragraph()
        sr = sig.add_run(
            "________________________________    "
            "________________________________    "
            "________________"
        )
        sr.font.size = Pt(9)
        sr.font.name = self._FONT

        lbl = doc.add_paragraph()
        lr = lbl.add_run("Signature" + " " * 36 + "Printed Name" + " " * 36 + "Date")
        lr.font.size = Pt(7)
        lr.font.name = self._FONT

    def _page_footer(self, doc: Document, ctx: dict) -> None:
        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            f"CONFIDENTIAL  \u00b7  {ctx['organization_name']}  \u00b7  "
            f"Compliance Period {ctx['reporting_period']}  \u00b7  "
            f"Generated via BC LCFS Portal  \u00b7  {ctx['generated_date']}"
        )
        fr.font.size = Pt(7)
        fr.font.name = self._FONT
